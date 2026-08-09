"""Regenerate the course-facing summary sheet and technical supplement.

The ML4H LaTeX paper is the source of truth. These DOCX files are compact
course-submission companions that restate the venue facts, GenAI disclosure,
and reproducibility snapshot without adding unverified claims.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SUMMARY_DOCX = ROOT / "summary" / "Module_14_Summary_Sheet_Pruned_Evidence_Gate_Ravi_Bajaj.docx"
SUPPLEMENT_DOCX = ROOT / "supplement" / "Module_14_Technical_Supplement_Pruned_Evidence_Gate_Ravi_Bajaj.docx"

TITLE = "A Deterministic Evidence Gate for Pharmacogenomic LLM Alerts: Synthetic Stress-Testing with Ablation"

ABSTRACT = (
    "LLM-generated medication alerts can sound more certain than their evidence permits. "
    "This capstone tests a deterministic evidence gate for pharmacogenomic alert drafts using "
    "33 author-designed synthetic cases: 23 overclaim archetypes and 10 bounded aligned alerts. "
    "The gate consumes structured annotations for source support, population fit, endpoint strength, "
    "and actionability, then returns allow, narrow, abstain, or deny. With oracle annotations, the "
    "full monitor routed 23/23 overclaims and preserved 10/10 bounded alerts; disabling source support, "
    "population fit, or claim strength let 2/23, 1/23, and 6/23 overclaims pass unchanged. Text-only "
    "extraction exposed the bottleneck: with citation fields, GPT-5.6-terra extracted 78/132 annotation "
    "fields and Grok-4.5 extracted 86/132; without citation fields, they extracted 73/132 and 70/132, "
    "while bounded-alert preservation dropped by 10 and 8 cases. These are specification-conformance "
    "and extraction-bottleneck results on synthetic cases, not clinical safety, generalization, or "
    "patient-benefit evidence."
)


def _set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def _set_cell_margins(table, top=80, start=120, bottom=80, end=120) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_cell_mar = tbl_pr.first_child_found_in("w:tblCellMar")
    if tbl_cell_mar is None:
        tbl_cell_mar = OxmlElement("w:tblCellMar")
        tbl_pr.append(tbl_cell_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tbl_cell_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tbl_cell_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_table_widths(table, widths: list[float]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for row in table.rows:
        for cell, width in zip(row.cells, widths, strict=True):
            cell.width = Inches(width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    _set_cell_margins(table)


def _style_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.08

    for name, size in (("Heading 1", 15), ("Heading 2", 12), ("Heading 3", 11)):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(46, 116, 181)
        style.paragraph_format.space_before = Pt(8 if name == "Heading 1" else 5)
        style.paragraph_format.space_after = Pt(4)


def _title(doc: Document, title: str, subtitle: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(title)
    run.bold = True
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(11, 37, 69)

    sub = doc.add_paragraph()
    sub.paragraph_format.space_after = Pt(8)
    r = sub.add_run(subtitle)
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(85, 85, 85)


def _kv_table(doc: Document, rows: list[tuple[str, str]], widths: list[float] | None = None):
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    widths = widths or [1.85, 5.15]
    _set_table_widths(table, widths)
    hdr = table.rows[0].cells
    hdr[0].text = "Field"
    hdr[1].text = "Value"
    for cell in hdr:
        _set_cell_shading(cell, "E8EEF5")
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
    for key, value in rows:
        cells = table.add_row().cells
        cells[0].text = key
        cells[1].text = value
        for p in cells[0].paragraphs:
            for r in p.runs:
                r.bold = True
    return table


def _matrix_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    _set_table_widths(table, widths)
    for cell, header in zip(table.rows[0].cells, headers, strict=True):
        cell.text = header
        _set_cell_shading(cell, "E8EEF5")
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
    for row in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, row, strict=True):
            cell.text = value
    return table


def _bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        p.add_run(item)


def build_summary() -> None:
    doc = Document()
    _style_doc(doc)
    _title(doc, "Summary Sheet", f"Ravi Bajaj | AI in Healthcare | Final paper: {TITLE}")

    doc.add_heading("Conference / Symposium", level=1)
    _kv_table(
        doc,
        [
            ("Name", "Machine Learning for Health Symposium (ML4H 2026)"),
            ("Conference URL", "https://ml4h.ahli.cc/"),
            ("Instructions / CFP", "https://ml4h.ahli.cc/submit/call-for-papers/"),
            ("Submission deadline", "September 10, 2026, 11:59 PM AoE"),
            ("Author notification", "October 22, 2026"),
            ("Camera-ready deadline", "November 7, 2026 (tentative)"),
            ("Event dates", "December 6-7, 2026"),
            ("Event location", "Sydney, Australia"),
            ("Suggested track", "Findings Track"),
            ("Editable paper format", "LaTeX/Overleaf source: paper/ml4h_findings_evidence_gate.tex and paper/ml4h_findings_refs.bib"),
        ],
    )

    doc.add_heading("Final Paper Title and Abstract", level=1)
    doc.add_paragraph(f"Title: {TITLE}")
    doc.add_paragraph(f"Abstract: {ABSTRACT}")

    doc.add_heading("Generative AI Use Disclosure", level=1)
    doc.add_paragraph(
        "Generative AI tools were used for brainstorming, critique, code scaffolding, drafting, formatting, "
        "and synthetic LLM experiments. GPT-5.6-terra and Grok-4.5 were used in LLM self-evaluation and "
        "text-only extraction experiments; GPT-5.6-sol generated unlabeled held-out scenarios for later "
        "author labeling. No real patient records, protected health information, or private genomic data were "
        "used or transmitted. The author selected the final scope, verified sources, ran the code, reviewed "
        "outputs, and remains responsible for all claims and limitations."
    )

    doc.add_heading("Optional Data, Code, and Tools", level=1)
    _bullets(
        doc,
        [
            "Synthetic 33-case pharmacogenomic/genomic medication-alert case bank.",
            "Deterministic three-check monitor with unit tests and replayable CSV/JSON outputs.",
            "Ablation, policy-comparator, LLM self-evaluation, text-only extraction, and held-out authoring artifacts under code/ and results/.",
            "Safety boundary: no diagnosis, no treatment recommendation, no real patient data, and no protected health information.",
        ],
    )
    SUMMARY_DOCX.parent.mkdir(exist_ok=True)
    doc.save(SUMMARY_DOCX)


def build_supplement() -> None:
    doc = Document()
    _style_doc(doc)
    _title(doc, "Technical Supplement", f"Reproducibility notes for {TITLE}")

    doc.add_heading("Safety Boundary", level=1)
    doc.add_paragraph(
        "Synthetic-only research artifact. No real patient data, protected health information, diagnosis, "
        "or treatment recommendation. The deterministic monitor evaluates claim permission levels from "
        "structured annotations; text-only extraction is evaluated separately and remains below the oracle "
        "annotation ceiling."
    )

    doc.add_heading("Reproducibility Commands", level=1)
    _bullets(
        doc,
        [
            r"py -3.13 code\pruned_evidence_gate.py",
            r"py -3.13 code\llm_self_evaluation_baseline.py",
            r"py -3.13 code\text_only_extraction.py",
            r"py -3.13 code\heldout_case_authoring.py",
            r"py -3.13 code\generate_paper_assets.py",
            r"py -3.13 -m pytest code -q",
            r"py -3.13 code\build_submission_package.py",
        ],
    )

    doc.add_heading("Main Synthetic Case Bank", level=1)
    doc.add_paragraph(
        "The final case matrix contains 33 author-designed synthetic cases: 10 bounded aligned alerts and "
        "23 designed overclaim archetypes. The 23/10 split is stress-test coverage, not a measured clinical base rate."
    )
    _matrix_table(
        doc,
        ["Metric", "Value"],
        [
            ["Case count", "33"],
            ["Author-designed overclaim cases", "23"],
            ["Bounded aligned alerts", "10"],
            ["Full monitor overclaims allowed unchanged", "0/23"],
            ["Bounded alerts preserved", "10/10"],
            ["Inappropriate denials", "0/10"],
            ["Action conformance", "33/33"],
            ["Primary-check conformance", "31/33"],
            ["Precedence-sensitive cases", "PGX19; PGX24"],
        ],
        [3.1, 3.9],
    )

    doc.add_page_break()
    doc.add_heading("Ablation and Comparator Snapshot", level=1)
    _matrix_table(
        doc,
        ["Analysis", "Result"],
        [
            ["Disable source support", "2/23 overclaims pass unchanged; 7 action changes; 12 primary-check changes."],
            ["Disable population fit", "1/23 overclaim passes unchanged; 5 action changes; 5 primary-check changes."],
            ["Disable claim strength", "6/23 overclaims pass unchanged; 6 action changes; 6 primary-check changes."],
            ["Ungated allow-all", "23/23 overclaims pass unchanged."],
            ["Claim-strength-only", "3/23 overclaims pass unchanged: PGX31, PGX32, PGX33."],
            ["Full three-check monitor", "0/23 overclaims pass unchanged."],
        ],
        [2.6, 4.4],
    )

    doc.add_heading("LLM Self-Evaluation Arms", level=1)
    _matrix_table(
        doc,
        ["Arm", "Scope", "Result"],
        [
            ["Arm A GPT-5.6-terra", "Full 33-case bank with source-support labels", "23/23 overclaims routed; 10/10 bounded alerts preserved; 31/33 action agreement; 25/33 primary-check agreement."],
            ["Arm A Grok-4.5", "Full 33-case bank with source-support labels", "23/23 overclaims routed; 10/10 bounded alerts preserved; 31/33 action agreement; 32/33 primary-check agreement."],
            ["Arm B GPT-5.6-terra", "PGX31-PGX33; source-support label withheld", "Missed PGX31; matched PGX32 and PGX33."],
            ["Arm B Grok-4.5", "PGX31-PGX33; source-support label withheld", "Matched PGX31 and PGX32; PGX33 action differed while primary check matched."],
            ["Uniform-NARROW sanity row", "Synthetic, no API call", "23/23 overclaims routed; 0/10 bounded alerts preserved; 33/33 narrowed."],
        ],
        [1.8, 2.25, 2.95],
    )

    doc.add_page_break()
    doc.add_heading("Text-Only Extraction", level=1)
    _matrix_table(
        doc,
        ["Provider / condition", "Field accuracy", "Downstream action", "Oracle gap"],
        [
            ["GPT-5.6-terra with citation", "78/132", "21/33", "Bounded preservation drop 6; inappropriate denial increase 0."],
            ["Grok-4.5 with citation", "86/132", "21/33", "Bounded preservation drop 1; inappropriate denial increase 0."],
            ["GPT-5.6-terra no citation", "73/132", "11/33", "Bounded preservation drop 10; inappropriate denial increase 10."],
            ["Grok-4.5 no citation", "70/132", "12/33", "Bounded preservation drop 8; inappropriate denial increase 8."],
        ],
        [2.15, 1.25, 1.5, 2.1],
    )
    doc.add_paragraph(
        "This is the main bottleneck result: the deterministic monitor performs as specified under oracle "
        "annotations, but LLM extraction from text and citation fields loses annotation accuracy and changes "
        "downstream actions. The no-citation control shows that missing citation fields shift the system toward "
        "over-denial of bounded alerts."
    )

    doc.add_heading("Held-Out Worksheet", level=1)
    doc.add_paragraph(
        "GPT-5.6-sol generated 12 model-authored held-out cases, author-labeled. They remain unlabeled in this "
        "package. The CSV worksheet has blank label columns, and no held-out accuracy is reported."
    )

    doc.add_heading("Traceability", level=1)
    doc.add_paragraph(
        "Every number in the paper is mapped to a generated file in results/paper_number_trace.csv. Raw model "
        "payloads are retained under results/*raw.json. Regression tests assert that author-only annotation_note "
        "fields are not sent in any model payload."
    )
    SUPPLEMENT_DOCX.parent.mkdir(exist_ok=True)
    doc.save(SUPPLEMENT_DOCX)


def main() -> None:
    build_summary()
    build_supplement()
    print(f"Wrote {SUMMARY_DOCX.relative_to(ROOT)}")
    print(f"Wrote {SUPPLEMENT_DOCX.relative_to(ROOT)}")


if __name__ == "__main__":
    main()
